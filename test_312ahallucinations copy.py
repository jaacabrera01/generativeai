import openai
import os
from docx import Document

openai.api_key = os.getenv("OPENAI_API_KEY")
doc = Document()
doc.add_heading('GenAI Hallucination Experiment in Software Testing', 0)

# Context: Only minimal info is given about the feature
feature_context = """
Feature: Secure File Upload
Description: Users can upload files to their account.
"""

# Ambiguous prompt (likely to cause hallucinations)
prompt_ambiguous = f"""Based on the following feature description, generate 5 acceptance criteria for testing.

{feature_context}
"""

# Specific prompt (discourages hallucinations)
prompt_specific = f"""Based on the following feature description, generate 5 acceptance criteria for testing. 
Only use information explicitly stated in the description. Do not invent or assume any additional requirements.

{feature_context}
"""

prompts = [("Ambiguous Prompt", prompt_ambiguous), ("Specific Prompt", prompt_specific)]
outputs = []

for label, prompt in prompts:
    print(f"\n--- {label} ---")
    print("Prompt:\n", prompt)
    response = openai.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}]
    )
    output = response.choices[0].message.content
    print("AI Output:\n", output)
    outputs.append((label, prompt, output))
    doc.add_heading(label, level=1)
    doc.add_paragraph(f"Prompt:\n{prompt}\n\nAI Output:\n{output}")

# Ask the LLM to analyze its own outputs for hallucinations
analysis_prompt = f"""Review the following two sets of acceptance criteria. For each, identify any criteria that are not supported by the provided feature description (i.e., hallucinations or invented requirements). List the hallucinated elements and explain why they are not justified by the context.

Ambiguous Prompt Output:
{outputs[0][2]}

Specific Prompt Output:
{outputs[1][2]}
"""

analysis_response = openai.chat.completions.create(
    model="gpt-4o",
    messages=[{"role": "user", "content": analysis_prompt}]
)
analysis = analysis_response.choices[0].message.content

print("\n--- Hallucination Analysis ---")
print(analysis)
doc.add_heading('Hallucination Analysis', level=1)
doc.add_paragraph(analysis)
doc.save("genai_hallucination_experiment_report.docx")
print("\nExperiment results and analysis have been saved to genai_hallucination_experiment_report.docx")