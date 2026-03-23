import openai
import os
from docx import Document

# Step 1: Create a new Word document
doc = Document()
doc.add_heading('Test Monitoring Metrics Report', 0)

# Step 2: Set up API key
openai.api_key = os.getenv("OPENAI_API_KEY")

test_data = """
Total Test Cases: 50
Executed: 40
Passed: 30
Failed: 10
Defects Logged: 8
Critical Defects: 3
"""

messages = [
    {
        "role": "system",
        "content": "You are a QA assistant that transforms test data into monitoring metrics and summaries."
    },
    {
        "role": "user",
        "content": f"""
Analyze the following test data and produce:
1. Key metrics (progress %, pass rate, defect density, coverage).
2. A simple dashboard table.
3. A natural language summary for stakeholders.

Test Data:
{test_data}
"""
    }
]

# Step 3: Send prompt to LLM
response = openai.chat.completions.create(
    model="gpt-4o",
    messages=messages
)

#Step 4: Print response to terminal
print("AI-Generated Test Monitoring Metrics:\n")
print(response.choices[0].message.content)

# Step 5: Save analysis to Word document
doc.add_heading('AI-Generated Test Monitoring Metrics', level=1)
doc.add_paragraph(response.choices[0].message.content)
doc.save("test_monitoring_metrics_report.docx")
print("\nAnalysis has also been saved to test_monitoring_metrics_report.docx")
