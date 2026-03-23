import openai
import os
from docx import Document


# Step 1: Create a new Word document
doc = Document()
doc.add_heading('Test Monitoring Metrics Report', 0)
# Step 2: Set up API key
openai.api_key = os.getenv("OPENAI_API_KEY")

# Step 3: Example regression test report (simplified)
test_report = """
Test Case 1: Login with valid credentials -> PASS
Test Case 2: Login with invalid password -> FAIL (Error message not displayed)
Test Case 3: Search functionality -> PASS
Test Case 4: Logout after login -> FAIL (Logout button not working)
Test Case 5: Profile update -> FAIL (Changes not saved)
"""

test_specification = """
System must:
1. Allow login with valid credentials.
2. Reject invalid passwords with error message.
3. Provide search functionality.
4. Allow logout after login.
5. Save profile updates correctly.
"""

# Step 4: Structured prompt for regression analysis
messages = [
    {
        "role": "system",
        "content": "You are a QA assistant that analyzes regression test reports step-by-step."
    },
    {
        "role": "user",
        "content": f"""
Structured Prompt Components:
Role: QA assistant
Context: Regression testing analysis
Instruction: Analyze the test report against the specification, cluster defects, maintain anomalies list, and cross-check findings.
Input Data: 
Test Report:
{test_report}

Test Specification:
{test_specification}

Constraints: Be methodical, link each step to the next, avoid vague answers.
Output Format: Provide a structured analysis with sections:
1. Comparison with specification
2. Defect clustering
3. Known anomalies list
4. Cross-check findings
5. Actionable insights
"""
    }
]

# Step 5: Send to LLM
response = openai.chat.completions.create(
    model="gpt-4o",
    messages=messages
)

# Step 6: Print structured analysis
print("Regression Test Report Analysis:\n")
print(response.choices[0].message.content)

# Step 7: Save analysis to Word document
doc.add_heading('AI-Generated Regression Test Report Analysis', level=1)
doc.add_paragraph(response.choices[0].message.content)
doc.save("regression_test_report_analysis.docx")
print("\nAnalysis has also been saved to regression_test_report_analysis.docx")
