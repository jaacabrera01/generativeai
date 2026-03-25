

# Step 1: Install necessary libraries
# !pip install transformers datasets torch
from transformers import AutoModelForCausalLM, AutoTokenizer, Trainer, TrainingArguments, DataCollatorForLanguageModeling, pipeline
from datasets import load_dataset
import docx

# Load a small language model and tokenizer
model_name = "distilgpt2"
model = AutoModelForCausalLM.from_pretrained(model_name)
tokenizer = AutoTokenizer.from_pretrained(model_name)
tokenizer.pad_token = tokenizer.eos_token

# Load your dataset using datasets library
dataset = load_dataset('text', data_files={'train': 'train.txt'})
train_dataset = dataset['train']

# Tokenize the dataset
def tokenize_function(examples):
    return tokenizer(examples["text"], truncation=True, padding="max_length", max_length=64)

# Tokenize and format the dataset for PyTorch
tokenized_dataset = train_dataset.map(tokenize_function, batched=True)
tokenized_dataset.set_format(type="torch", columns=["input_ids", "attention_mask"])

# Set up data collator and training arguments
data_collator = DataCollatorForLanguageModeling(
    tokenizer=tokenizer, mlm=False,
)
training_args = TrainingArguments(
    output_dir="./slm-finetuned",
    num_train_epochs=3,
    per_device_train_batch_size=2,
    save_steps=10_000,
    save_total_limit=2,
    logging_steps=10,
)

# Initialize the Trainer and start fine-tuning
trainer = Trainer(
    model=model,
    args=training_args,
    data_collator=data_collator,
    train_dataset=tokenized_dataset,
)
print("Starting fine-tuning...")
trainer.train()
print("Fine-tuning complete.")

# Save the fine-tuned model and tokenizers
model.save_pretrained("./slm-finetuned")
tokenizer.save_pretrained("./slm-finetuned")
print("Model saved to ./slm-finetuned")

# Test the fine-tuned model
generator = pipeline('text-generation', model='./slm-finetuned')
prompt = "Login failed: account locked."
result = generator(prompt, max_length=50)
print("\nTest output for prompt:")
print(result[0]['generated_text'])

# Optionally, save the result to a Word document
doc = docx.Document()
doc.add_heading('GenAI Fine-Tuning Experiment', 0)
doc.add_heading('Prompt', level=1)
doc.add_paragraph(prompt)
doc.add_heading('Generated Output', level=1)
doc.add_paragraph(result[0]['generated_text'])
doc.save("GenAI_FineTuning_Experiment.docx")
print("\nDocument saved as GenAI_FineTuning_Experiment.docx")