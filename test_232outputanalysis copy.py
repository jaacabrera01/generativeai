import openai
import os
from docx import Document

# Set your OpenAI API key
openai.api_key = os.getenv("OPENAI_API_KEY")

# Create a new Word document
doc = Document()
doc.add_heading('A/B Prompt Testing Report', 0)

# Define two prompts for A/B testing
prompt_A = "Generate three test cases for the login feature of a banking app."
prompt_B = "Generate three keyword-driven test cases for the login feature of a banking app. Each test case should include: Title, Steps, and Expected Result."

prompts = [("Prompt A", prompt_A), ("Prompt B", prompt_B)]
outputs = []

# Run both prompts and collect outputs
for label, prompt in prompts:
    print(f"\n--- {label} ---")
    print("Prompt:\n", prompt)
    response = openai.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}]
    )
    output = response.choices[0].message.content
    print("AI Output:\n", output)
    outputs.append((label, prompt, output))

    # Save to Word document
    doc.add_heading(label, level=1)
    doc.add_paragraph(f"Prompt:\n{prompt}\n\nAI Output:\n{output}")

# Output analysis: ask the LLM to compare the results
analysis_prompt = f"""
Compare the following two outputs for the same test task. Score each for correctness, coverage, and clarity. Suggest which prompt is better and why.

Prompt A Output:
{outputs[0][2]}

Prompt B Output:
{outputs[1][2]}
"""

analysis_response = openai.chat.completions.create(
    model="gpt-4o",
    messages=[{"role": "user", "content": analysis_prompt}]
)
analysis = analysis_response.choices[0].message.content

print("\n--- Output Analysis ---")
print(analysis)

# Save analysis to Word
doc.add_heading('AI Output Analysis', level=1)
doc.add_paragraph(analysis)
doc.save("ab_prompt_testing_report.docx")
print("\nA/B test results and analysis have also been saved to ab_prompt_testing_report.docx")