import openai
import os
import base64
from docx import Document

#Step 1: Create a new Word document
doc = Document()
doc.add_heading('Iterative Prompting Test Report', 0)

#Step 2: Set up API key
openai.api_key = os.getenv("OPENAI_API_KEY")

#Step 3: List of prompt versions (iterate and refine as needed)
prompts = [
    # Initial prompt (simple)
    "Generate three test cases for the login feature of a banking app.",

    # First refinement: add structure
    "Generate three test cases for the login feature of a banking app. Each test case should include: Title, Steps, and Expected Result.",

    # Second refinement: specify positive and negative cases
    "Generate three test cases for the login feature of a banking app. Include at least one positive and one negative test. Each test case should have: Title, Steps, and Expected Result.",

    # Third refinement: require traceability to requirements
    "Generate three test cases for the login feature of a banking app. Each test case should include: Title, Steps, Expected Result, and the requirement it covers. Include both positive and negative scenarios."
]

for i, prompt in enumerate(prompts, 1):
    print(f"\n--- Prompt Version {i} ---")
    print("Prompt:\n", prompt)
    response = openai.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}]
    )
    print("AI Output:\n", response.choices[0].message.content)

# Step 4: Save final iteration's output to Word document
final_output = response.choices[0].message.content
doc.add_heading('Final Iteration Output', level=1)
doc.add_paragraph(final_output)
doc.save("iterative_prompting_test_report.docx")
print("\nFinal output has also been saved to iterative_prompting_test_report.docx") 

#Step 5: Optionally, print to terminal
print("\nFinal Iteration Output:\n")
print(final_output) 