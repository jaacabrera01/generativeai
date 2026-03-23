import openai
import os
import base64
from docx import Document

# Step 1: Create a new Word document
doc = Document()
doc.add_heading('User Feedback Analysis Report', 0) 

# Step 2: Set up API key
openai.api_key = os.getenv("OPENAI_API_KEY")

# Step 3: Sample user feedback data
user_feedback = """
1. "The app crashes every time I try to upload a photo."
2. "I love the new design, it's very intuitive!"
3. "The search function is really slow and often doesn't return relevant results."
4. "Customer support was very helpful when I had an issue with my account."
5. "I wish there were more customization options for notifications."
"""
# Step 4: Create a prompt for analyzing user feedback
# Short/general prompt
prompt_short = f"Categorize each of the following user feedback comments:\n{user_feedback}"

# Long/specific prompt
prompt_long = f"""Analyze the following user feedback and categorize each comment into one of the following categories: Bug Report, Positive Feedback, Feature Request, or Other. For each comment, provide the category and a brief summary. If a comment fits more than one category, choose the most relevant one.
User Feedback:
{user_feedback}
"""

prompts = [("Short Prompt", prompt_short), ("Long/Specific Prompt", prompt_long)]
outputs = []

for label, prompt in prompts:
    print(f"\n--- {label} ---")
    print("Prompt:\n", prompt)
    response = openai.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}]
    )
    output = response.choices[0].message.content
    print("AI Output:\n", output)
    outputs.append((label, prompt, output))

    # Save to Word document
    doc.add_heading(label, level=1)
    doc.add_paragraph(f"Prompt:\n{prompt}\n\nAI Output:\n{output}")

# (Optional) Ask the LLM to compare the outputs
compare_prompt = f"""Compare the following two analyses of user feedback. Which is more useful and why? Consider accuracy, clarity, and usefulness for a product team.

Short Prompt Output:
{outputs[0][2]}

Long/Specific Prompt Output:
{outputs[1][2]}
"""

compare_response = openai.chat.completions.create(
    model="gpt-4o",
    messages=[{"role": "user", "content": compare_prompt}]
)
comparison = compare_response.choices[0].message.content

print("\n--- Output Comparison ---")
print(comparison)

doc.add_heading('Output Comparison', level=1)
doc.add_paragraph(comparison)
doc.save("prompt_length_experiment_report.docx")
print("\nExperiment results and analysis have been saved to prompt_length_experiment_report.docx")