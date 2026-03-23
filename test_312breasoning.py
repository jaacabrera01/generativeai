import openai
import os
from docx import Document

openai.api_key = os.getenv("OPENAI_API_KEY")
doc = Document()
doc.add_heading('GenAI Reasoning Error Experiment in Test Planning', 0)

# Test planning scenario: Test case prioritization
scenario = """
You have 5 test cases for a login feature:
1. Valid login with correct credentials
2. Invalid login with wrong password
3. Invalid login with empty fields
4. Password reset flow
5. Account lockout after 5 failed attempts

Constraints:
- The password reset flow is currently broken and cannot be tested.
- Account lockout is a high-risk area due to recent security incidents.
- Valid login is the most frequently used path.
- You have time to run only 3 test cases.

Task: Prioritize the 5 test cases and select the top 3 to execute, justifying your choices.
"""

# The correct answer (for comparison)
correct_answer = """
1. Account lockout after 5 failed attempts (high risk)
2. Valid login with correct credentials (most used)
3. Invalid login with wrong password (common error path)
(Password reset is excluded as it cannot be tested; empty fields is less critical)
"""

# Prompt variations
prompt_basic = f"""Prioritize the following 5 test cases for a login feature and select the top 3 to execute, justifying your choices. Consider the constraints and risks described.

{scenario}
"""

prompt_step_by_step = f"""Let's solve this step by step. First, list the constraints and risks. Then, prioritize the 5 test cases for a login feature and select the top 3 to execute, justifying your choices.

{scenario}
"""

prompts = [("Basic Prompt", prompt_basic), ("Step-by-Step Prompt", prompt_step_by_step)]
models = [
    ("GPT-4o (LLM)", "gpt-4o"),
    ("GPT-3.5-turbo (SLM)", "gpt-3.5-turbo"),
    # Add another reasoning model if available, e.g., "claude-3-opus" or similar
]

results = []

for model_label, model_name in models:
    for prompt_label, prompt in prompts:
        print(f"\n--- {model_label} | {prompt_label} ---")
        print("Prompt:\n", prompt)
        response = openai.chat.completions.create(
            model=model_name,
            messages=[{"role": "user", "content": prompt}]
        )
        output = response.choices[0].message.content
        print("AI Output:\n", output)
        results.append((model_label, prompt_label, prompt, output))
        doc.add_heading(f"{model_label} - {prompt_label}", level=1)
        doc.add_paragraph(f"Prompt:\n{prompt}\n\nAI Output:\n{output}")

# Compare each output to the correct answer
for model_label, prompt_label, prompt, output in results:
    analysis_prompt = f"""Compare the following AI output to the correct answer for the test case prioritization task. Identify any reasoning errors, such as ignoring constraints, misjudging risk, or failing to justify choices.

AI Output:
{output}

Correct Answer:
{correct_answer}
"""
    analysis_response = openai.chat.completions.create(
        model="gpt-4o",  # Use your best model for analysis
        messages=[{"role": "user", "content": analysis_prompt}]
    )
    analysis = analysis_response.choices[0].message.content
    print(f"\n--- Analysis for {model_label} | {prompt_label} ---")
    print(analysis)
    doc.add_heading(f"Analysis: {model_label} - {prompt_label}", level=2)
    doc.add_paragraph(analysis)

doc.save("genai_reasoning_error_experiment_report.docx")
print("\nExperiment results and analysis have been saved to genai_reasoning_error_experiment_report.docx")