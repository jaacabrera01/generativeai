import openai
import os
from docx import Document

# Set your OpenAI API key
openai.api_key = os.getenv("OPENAI_API_KEY")

# Create a new Word document
doc = Document()
doc.add_heading('Test Monitoring Metrics Report', 0)

# Example test data (replace with your real data)
test_data = """
Test Results:
- Total test cases: 100
- Passed: 80
- Failed: 15
- Blocked: 5

Defect Log:
- New defects this cycle: 7
- Open defects: 12
- Critical defects: 2

Coverage:
- Requirements covered: 90/100
- Uncovered requirements: 10
"""

prompt = f"""
You are an AI test analyst. Given the following test data, generate key test monitoring metrics such as test progress, defect trends, and coverage. Highlight any potential risks. Then, summarize the findings in simple language for stakeholders.

Test Data:
{test_data}
"""

response = openai.chat.completions.create(
    model="gpt-4o",
    messages=[{"role": "user", "content": prompt}]
)

metrics_summary = response.choices[0].message.content

# Print to terminal
print("AI-Generated Test Monitoring Metrics and Summary:\n")
print(metrics_summary)

# Save to Word file
doc.add_heading('AI-Generated Test Monitoring Metrics and Summary', level=1)
doc.add_paragraph(metrics_summary)
doc.save("test_monitoring_metrics_report.docx")
print("\nAll outputs have also been saved to test_monitoring_metrics_report.docx")