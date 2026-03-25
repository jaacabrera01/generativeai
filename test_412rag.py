import openai
import os
from docx import Document

print("OPENAI_API_KEY:", os.getenv("OPENAI_API_KEY"))

#Step 1 Create word document and add heading
doc = Document()
doc.add_heading('GenAI RAG Experiment in Test Planning', 0)

#Step 2 Set up OpenAI API key and define prompts
openai.api_key = os.getenv("OPENAI_API_KEY")

# Baseline prompt
prompt = "Explain the process of login error analysis in software testing."

# Simulated retrieval (RAG)
retrieved_docs = [
    "In software testing, login errors are often caused by incorrect credentials, network issues, or server misconfigurations.",
    "Best practices include anonymizing user data and logging access for audit trails."
]
rag_prompt = "\n".join(retrieved_docs) + "\n\n" + prompt

def get_llm_output(prompt):
    response = openai.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content

# Get LLM outputs and print
no_rag_output = get_llm_output(prompt)
rag_output = get_llm_output(rag_prompt)

print("--- LLM Output (No RAG) ---")
print(no_rag_output)

print("\n--- LLM Output (With RAG) ---")
print(rag_output)

# Add outputs to the document
doc.add_heading('LLM Output (No RAG)', level=1)
doc.add_paragraph(no_rag_output)

doc.add_heading('LLM Output (With RAG)', level=1)
doc.add_paragraph(rag_output)

doc.save("GenAI_RAG_Experiment.docx")
print("\nDocument saved as GenAI_RAG_Experiment.docx")