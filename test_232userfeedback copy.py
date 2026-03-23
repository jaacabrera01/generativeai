import openai
import os
import base64
from docx import Document

# Step 1: Create a new Word document
doc = Document()
doc.add_heading('User Feedback Analysis Report', 0) 

# Step 2: Set up API key
openai.api_key = os.getenv("OPENAI_API_KEY")

# Step 3: Sample user feedback data
user_feedback = """
1. "The app crashes every time I try to upload a photo."
2. "I love the new design, it's very intuitive!"
3. "The search function is really slow and often doesn't return relevant results."
4. "Customer support was very helpful when I had an issue with my account."
5. "I wish there were more customization options for notifications."
"""         
# Step 4: Create a prompt for analyzing user feedback
prompt = f"""Analyze the following user feedback and categorize each comment into one of the following categories: Bug Report, Positive Feedback, Feature Request, or Other. Provide a brief summary for each comment.
User Feedback:
{user_feedback}
""" 
# Step 5: Send prompt to LLM
response = openai.chat.completions.create(
    model="gpt-4o",
    messages=[{"role": "user", "content": prompt}]
)
analysis = response.choices[0].message.content
print("\n--- User Feedback Analysis ---")
print(analysis)

# Step 6: Save analysis to Word document
doc.add_heading('User Feedback Analysis', level=1)
doc.add_paragraph(analysis)
doc.save("user_feedback_analysis_report.docx")
print("\nUser feedback analysis has also been saved to user_feedback_analysis_report.docx")
