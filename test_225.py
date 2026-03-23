import openai
import os
from docx import Document

# Step 1: Create a new Word document
doc = Document()
doc.add_heading('Test Monitoring Metrics Report', 0)

# Step 2: Set up API key
openai.api_key = os.getenv("OPENAI_API_KEY")

#Step 3: Example regression test report (simplified)
test_report = """
Test Case 1: Login with valid credentials -> PASS
Test Case 2: Login with invalid password -> FAIL (Error message not displayed)
Test Case 3: Search functionality -> PASS
Test Case 4: Logout after login -> FAIL (Logout button not working)
Test Case 5: Profile update -> FAIL (Changes not saved)
"""
print("\n--- Prompt Chaining Example ---")

# Step 4: List features to be tested
prompt1 = "List three important features to test in an online banking app."
response1 = openai.chat.completions.create(
    model="gpt-4o",
    messages=[{"role": "user", "content": prompt1}]
)
features = response1.choices[0].message.content
print("Prompt 1:\n", prompt1)
print("Response 1:\n", features)

# Step 5: Overwrite features with your own list (manual override)
features = """
1. User Login
2. Fund Transfer
3. Account Balance Inquiry
"""

# Step 6: Generate a test case for each feature
prompt2 = f"Generate a keyword-driven test case for each of these features:\n{features}"
response2 = openai.chat.completions.create(
    model="gpt-4o",
    messages=[{"role": "user", "content": prompt2}]
)
test_cases = response2.choices[0].message.content
print("\nPrompt 2:\n", prompt2)
print("Response 2:\n", test_cases)

# Step 7: Few-Shot Prompting Example
print("\n--- Few-Shot Prompting Example ---")

few_shot_prompt = """
Example 1:
Test Case: Login with valid credentials
Steps: 1. Open app 2. Enter username and password 3. Click Login
Expected Result: User is logged in

Example 2:
Test Case: Login with invalid password
Steps: 1. Open app 2. Enter username and wrong password 3. Click Login
Expected Result: Error message is shown

Now, generate a test case for "Password reset".
"""
response3 = openai.chat.completions.create(
    model="gpt-4o",
    messages=[{"role": "user", "content": few_shot_prompt}]
)
print("Few-Shot Prompt:\n", few_shot_prompt)
print("Few-Shot Response:\n", response3.choices[0].message.content)

# Step 8: Meta Prompting Example
print("\n--- Meta Prompting Example ---")

meta_prompt = """
I want to generate structured test cases for a web application using an LLM. Suggest an effective prompt I can use to get high-quality, consistent test cases in a table format.
"""
response4 = openai.chat.completions.create(
    model="gpt-4o",
    messages=[{"role": "user", "content": meta_prompt}]
)
print("Meta Prompt:\n", meta_prompt)
print("Meta Prompt Response:\n", response4.choices[0].message.content)

# Step 9: Save analysis to Word document
doc.add_heading('Prompt Chaining Test Cases', level=1)
doc.add_paragraph(test_cases)

doc.add_heading('Few-Shot Prompting Example', level=1)
doc.add_paragraph(response3.choices[0].message.content)

doc.add_heading('Meta Prompting Example', level=1)
doc.add_paragraph(response4.choices[0].message.content)

doc.save("prompt_test_report_analysis.docx")
print("\nAnalysis has also been saved to prompt_test_report_analysis.docx")