import docx 
# Example for OpenAI GPT-4o
num_calls = 1000
tokens_per_call = 1000
input_price_per_million = 5.0
output_price_per_million = 15.0

total_input_tokens = num_calls * (tokens_per_call // 2)
total_output_tokens = num_calls * (tokens_per_call // 2)

input_cost = (total_input_tokens / 1_000_000) * input_price_per_million
output_cost = (total_output_tokens / 1_000_000) * output_price_per_million
total_cost = input_cost + output_cost

print(f"Total input tokens: {total_input_tokens}")
print(f"Total output tokens: {total_output_tokens}")
print(f"Input cost: ${input_cost:.2f}")
print(f"Output cost: ${output_cost:.2f}")
print(f"Monthly total cost: ${total_cost:.2f}")
print("\nNote: Actual costs may vary based on the specific model used and any additional features or optimizations applied.")   
print("For more accurate cost estimation, refer to the OpenAI pricing page: https://openai.com/pricing")    
print("Consider implementing cost-saving strategies such as token optimization, batching requests, or using lower-cost models for less critical tasks.")    
print("Always monitor your usage and set up budget alerts to avoid unexpected costs.")  
print("Remember that fine-tuning and inference costs can add up, so it's important to plan and optimize your GenAI usage accordingly.") 
print("For more information on cost management and optimization, refer to OpenAI's best practices: https://openai.com/docs/best-practices/cost-management") 
print("In summary, while GenAI offers powerful capabilities, it's crucial to be mindful of the associated costs and implement strategies to manage and optimize your usage effectively.")   
print("By understanding the cost structure and implementing best practices, you can leverage GenAI's capabilities while keeping costs under control.")  
print("Always stay informed about the latest updates to pricing and cost management strategies to ensure you're making the most of your GenAI investment.") 

#Save to Word document
doc = docx.Document()
doc.add_heading('GenAI Cost Estimation', 0)
doc.add_paragraph(f"Total input tokens: {total_input_tokens}")
doc.add_paragraph(f"Total output tokens: {total_output_tokens}")
doc.add_paragraph(f"Input cost: ${input_cost:.2f}")
doc.add_paragraph(f"Output cost: ${output_cost:.2f}")
doc.add_paragraph(f"Monthly total cost: ${total_cost:.2f}")
doc.add_paragraph("\nNote: Actual costs may vary based on the specific model used and any additional features or optimizations applied.")   
doc.add_paragraph("For more accurate cost estimation, refer to the OpenAI pricing page: https://openai.com/pricing")    
doc.add_paragraph("Consider implementing cost-saving strategies such as token optimization, batching requests, or using lower-cost models for less critical tasks.")    
doc.add_paragraph("Always monitor your usage and set up budget alerts to avoid unexpected costs.")  
doc.add_paragraph("Remember that fine-tuning and inference costs can add up, so it's important to plan and optimize your GenAI usage accordingly.") 
doc.add_paragraph("For more information on cost management and optimization, refer to OpenAI's best practices: https://openai.com/docs/best-practices/cost-management") 
doc.add_paragraph("In summary, while GenAI offers powerful capabilities, it's crucial to be mindful of the associated costs and implement strategies to manage and optimize your usage effectively.")   
doc.add_paragraph("By understanding the cost structure and implementing best practices, you can leverage GenAI's capabilities while keeping costs under control.")
doc.add_paragraph("Always stay informed about the latest updates to pricing and cost management strategies to ensure you're making the most of your GenAI investment.") 
doc.save("GenAI_Cost_Estimation_Report.docx")
print("\nReport saved as GenAI_Cost_Estimation_Report.docx")